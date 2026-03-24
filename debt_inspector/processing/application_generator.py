"""
Генератор DOCX-документов для процедуры банкротства.

Генерирует:
- Заявление в арбитражный суд (ст. 213.4 ФЗ-127)
- Заявление в МФЦ (внесудебное банкротство)
- Список кредиторов
- Опись имущества
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _fmt_money(value: float) -> str:
    """Форматирует число как деньги."""
    return f"{value:,.2f}".replace(",", " ")


def _add_paragraph(doc: Document, text: str, bold: bool = False,
                   alignment=None, font_size: int = 12, space_after: int = 6):
    """Добавляет параграф с настройками."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "Times New Roman"
    run.bold = bold
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def generate_court_application(pipeline_data: dict, output_path: Path) -> Path:
    """Генерирует заявление в арбитражный суд о признании банкротом."""
    doc = Document()

    # Настройка стиля
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    d = pipeline_data
    full_name = " ".join(p for p in [d.get("last_name", ""), d.get("first_name", ""), d.get("middle_name", "")] if p)

    # Шапка — суд
    _add_paragraph(doc, f"В {d.get('court_name', 'Арбитражный суд')}", alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(doc, "", space_after=2)
    _add_paragraph(doc, "Заявитель (должник):", bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(doc, full_name, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    if d.get("birth_date"):
        _add_paragraph(doc, f"Дата рождения: {d['birth_date']}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)
    if d.get("inn"):
        _add_paragraph(doc, f"ИНН: {d['inn']}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)
    if d.get("snils"):
        _add_paragraph(doc, f"СНИЛС: {d['snils']}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)
    if d.get("passport_series"):
        _add_paragraph(
            doc,
            f"Паспорт: {d['passport_series']} {d.get('passport_number', '')}",
            alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11,
        )
        if d.get("passport_issued_by"):
            _add_paragraph(doc, f"выдан: {d['passport_issued_by']}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=10)
        if d.get("passport_issued_date"):
            _add_paragraph(doc, f"дата выдачи: {d['passport_issued_date']}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=10)

    address_parts = [d.get("region", ""), d.get("city", ""), d.get("address", "")]
    address = ", ".join(p for p in address_parts if p)
    if address:
        _add_paragraph(doc, f"Адрес: {address}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)
    if d.get("phone"):
        _add_paragraph(doc, f"Тел.: {d['phone']}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)
    if d.get("email"):
        _add_paragraph(doc, f"E-mail: {d['email']}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)

    # Заголовок
    _add_paragraph(doc, "", space_after=12)
    _add_paragraph(doc, "ЗАЯВЛЕНИЕ", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)
    _add_paragraph(doc, "о признании гражданина банкротом", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=13)
    _add_paragraph(doc, "", space_after=6)

    # Текст заявления
    total_debt = d.get("total_debt", 0)
    _add_paragraph(
        doc,
        f"В соответствии со ст. 213.4 Федерального закона от 26.10.2002 N 127-ФЗ "
        f"\"О несостоятельности (банкротстве)\" прошу признать меня, {full_name}, банкротом.",
    )
    _add_paragraph(
        doc,
        f"Общая сумма задолженности перед кредиторами составляет {_fmt_money(total_debt)} руб. "
        f"Обязательства не исполнены в течение более трёх месяцев с даты, "
        f"когда они должны были быть исполнены.",
    )
    _add_paragraph(
        doc,
        "Удовлетворение требований одного или нескольких кредиторов приводит к невозможности "
        "исполнения денежных обязательств в полном объёме перед другими кредиторами.",
    )

    # Неизвестные кредиторы
    if d.get("has_unknown_creditors"):
        total_estimated = d.get("total_estimated_debt", 0)
        _add_paragraph(doc, "", space_after=4)
        if total_estimated and total_estimated > total_debt:
            _add_paragraph(
                doc,
                f"Общая сумма обязательств оценивается мною приблизительно в "
                f"{_fmt_money(total_estimated)} руб. Помимо указанных выше кредиторов, "
                f"у меня имеются иные неисполненные обязательства, точные данные по которым "
                f"(наименования кредиторов, суммы задолженности, реквизиты договоров) "
                f"мне не известны в полном объёме.",
            )
        else:
            _add_paragraph(
                doc,
                "Помимо указанных выше кредиторов, у меня имеются иные неисполненные "
                "обязательства, точные данные по которым мне не известны в полном объёме.",
            )
        _add_paragraph(
            doc,
            "Прошу суд обязать финансового управляющего установить полный перечень "
            "кредиторов на основании данных бюро кредитных историй, ФССП, ФНС и "
            "иных источников информации.",
        )
        note = d.get("unknown_creditors_note", "")
        if note:
            _add_paragraph(doc, f"Дополнительно сообщаю: {note}", font_size=11)

    # Имущество
    properties = d.get("properties", [])
    has_property = any(p.get("property_type") for p in properties)
    if has_property:
        _add_paragraph(doc, "Имущество:", bold=True, space_after=4)
        for p in properties:
            if p.get("property_type"):
                desc = p.get("description", p["property_type"])
                val = _fmt_money(p.get("estimated_value", 0))
                sole = " (единственное жильё)" if p.get("is_sole_housing") else ""
                _add_paragraph(doc, f"- {desc} — {val} руб.{sole}", font_size=11, space_after=2)

    # Доходы
    emp = d.get("employment_type", "")
    income = d.get("monthly_income", 0)
    if emp == "employed":
        employer = d.get("employer", "")
        _add_paragraph(doc, f"Место работы: {employer or 'указано в приложении'}. Ежемесячный доход: {_fmt_money(income)} руб.")
    elif emp == "pensioner":
        _add_paragraph(doc, f"Являюсь пенсионером. Ежемесячный доход (пенсия): {_fmt_money(income)} руб.")
    elif emp == "unemployed":
        _add_paragraph(doc, "В настоящее время не работаю.")

    deps = d.get("dependents_count", 0)
    if deps > 0:
        _add_paragraph(doc, f"На иждивении: {deps} чел.")

    # Таблица кредиторов
    _add_paragraph(doc, "", space_after=4)
    _add_paragraph(doc, "Список кредиторов:", bold=True, space_after=4)

    creditors = d.get("creditors", [])
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    for i, text in enumerate(["N", "Кредитор", "Сумма, руб.", "Основание"]):
        headers[i].text = text
        for p in headers[i].paragraphs:
            p.runs[0].bold = True if p.runs else None
            p.runs[0].font.size = Pt(10) if p.runs else None

    for idx, c in enumerate(creditors, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        name = c.get("name", "")
        if c.get("creditor_inn"):
            name += f" (ИНН {c['creditor_inn']})"
        row[1].text = name
        row[2].text = _fmt_money(c.get("amount", 0))
        row[3].text = c.get("contract_number", "") or c.get("debt_type", "")

    # Итого
    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = "ИТОГО:"
    total_row[2].text = _fmt_money(total_debt)
    total_row[3].text = ""
    for cell in total_row:
        for p in cell.paragraphs:
            if p.runs:
                p.runs[0].bold = True

    # Просительная часть
    _add_paragraph(doc, "", space_after=8)
    _add_paragraph(doc, "На основании изложенного, руководствуясь ст. 213.4 ФЗ-127,", bold=True)
    _add_paragraph(doc, "ПРОШУ:", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"1. Признать меня, {full_name}, несостоятельным (банкротом).", space_after=2)
    _add_paragraph(
        doc,
        "2. Утвердить финансового управляющего из числа членов "
        "саморегулируемой организации арбитражных управляющих.",
    )

    # Приложения
    _add_paragraph(doc, "", space_after=8)
    _add_paragraph(doc, "Приложения:", bold=True, space_after=4)
    attachments = [
        "Копия паспорта", "Копия ИНН", "Копия СНИЛС",
        "Список кредиторов и должников", "Опись имущества",
        "Справки о доходах за 3 года", "Выписка из ЕГРН",
        "Справка ГИБДД", "Выписки из банков",
        "Копии кредитных договоров",
        "Квитанция об уплате госпошлины (300 руб.)",
        "Квитанция о внесении депозита (25 000 руб.)",
    ]
    for i, att in enumerate(attachments, 1):
        _add_paragraph(doc, f"{i}. {att}", font_size=11, space_after=1)

    # Подпись
    _add_paragraph(doc, "", space_after=16)
    _add_paragraph(doc, 'Дата: "____" _____________ 20___ г.')
    _add_paragraph(doc, "", space_after=8)
    _add_paragraph(doc, f"Подпись: _________________ / {full_name} /")

    doc.save(str(output_path))
    return output_path


def generate_mfc_application(pipeline_data: dict, output_path: Path) -> Path:
    """Генерирует заявление в МФЦ о внесудебном банкротстве."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    d = pipeline_data
    full_name = " ".join(p for p in [d.get("last_name", ""), d.get("first_name", ""), d.get("middle_name", "")] if p)

    # Шапка
    _add_paragraph(doc, "В МФЦ по месту жительства", alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(doc, "", space_after=2)
    _add_paragraph(doc, "Заявитель:", bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(doc, full_name, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    if d.get("birth_date"):
        _add_paragraph(doc, f"Дата рождения: {d['birth_date']}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)
    if d.get("inn"):
        _add_paragraph(doc, f"ИНН: {d['inn']}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)
    if d.get("snils"):
        _add_paragraph(doc, f"СНИЛС: {d['snils']}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)
    if d.get("passport_series"):
        _add_paragraph(doc, f"Паспорт: {d['passport_series']} {d.get('passport_number', '')}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)

    address_parts = [d.get("region", ""), d.get("city", ""), d.get("address", "")]
    address = ", ".join(p for p in address_parts if p)
    if address:
        _add_paragraph(doc, f"Адрес: {address}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)
    if d.get("phone"):
        _add_paragraph(doc, f"Тел.: {d['phone']}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)

    # Заголовок
    _add_paragraph(doc, "", space_after=12)
    _add_paragraph(doc, "ЗАЯВЛЕНИЕ", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)
    _add_paragraph(doc, "о признании гражданина банкротом во внесудебном порядке",
                   bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=13)
    _add_paragraph(doc, "", space_after=6)

    total_debt = d.get("total_debt", 0)
    _add_paragraph(
        doc,
        f"В соответствии со ст. 223.2 Федерального закона от 26.10.2002 N 127-ФЗ "
        f"\"О несостоятельности (банкротстве)\" прошу признать меня, {full_name}, "
        f"банкротом во внесудебном порядке.",
    )
    _add_paragraph(doc, f"Общая сумма обязательств составляет {_fmt_money(total_debt)} руб.")

    # Таблица кредиторов
    _add_paragraph(doc, "Список кредиторов:", bold=True, space_after=4)
    creditors = d.get("creditors", [])
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    for i, text in enumerate(["N", "Кредитор", "Сумма, руб."]):
        headers[i].text = text
        for p in headers[i].paragraphs:
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(10)

    for idx, c in enumerate(creditors, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = c.get("name", "")
        row[2].text = _fmt_money(c.get("amount", 0))

    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = "ИТОГО:"
    total_row[2].text = _fmt_money(total_debt)

    # Подпись
    _add_paragraph(doc, "", space_after=16)
    _add_paragraph(doc, 'Дата: "____" _____________ 20___ г.')
    _add_paragraph(doc, "", space_after=8)
    _add_paragraph(doc, f"Подпись: _________________ / {full_name} /")

    doc.save(str(output_path))
    return output_path


def generate_creditors_list(pipeline_data: dict, output_path: Path) -> Path:
    """Генерирует список кредиторов в DOCX."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)

    d = pipeline_data
    full_name = " ".join(p for p in [d.get("last_name", ""), d.get("first_name", ""), d.get("middle_name", "")] if p)

    _add_paragraph(doc, "СПИСОК КРЕДИТОРОВ И ДОЛЖНИКОВ", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)
    _add_paragraph(doc, f"гражданина {full_name}", alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=12)
    _add_paragraph(doc, "", space_after=8)

    creditors = d.get("creditors", [])
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    for i, text in enumerate(["N", "Наименование кредитора", "ИНН кредитора", "Сумма долга, руб.", "Основание"]):
        headers[i].text = text
        for p in headers[i].paragraphs:
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(9)

    total = 0
    for idx, c in enumerate(creditors, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = c.get("name", "")
        row[2].text = c.get("creditor_inn", "")
        amt = c.get("amount", 0)
        row[3].text = _fmt_money(amt)
        row[4].text = c.get("contract_number", "") or c.get("debt_type", "")
        total += amt

    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = "ИТОГО:"
    total_row[2].text = ""
    total_row[3].text = _fmt_money(total)
    total_row[4].text = ""
    for cell in total_row:
        for p in cell.paragraphs:
            if p.runs:
                p.runs[0].bold = True

    _add_paragraph(doc, "", space_after=12)
    _add_paragraph(doc, 'Дата: "____" _____________ 20___ г.')
    _add_paragraph(doc, f"Подпись: _________________ / {full_name} /")

    doc.save(str(output_path))
    return output_path


def generate_skip_restructuring_petition(pipeline_data: dict, output_path: Path) -> Path:
    """Генерирует ходатайство о введении процедуры реализации имущества,
    минуя процедуру реструктуризации долгов (ст. 213.6 п.8 ФЗ-127)."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    d = pipeline_data
    full_name = " ".join(
        p for p in [d.get("last_name", ""), d.get("first_name", ""), d.get("middle_name", "")] if p
    )

    # Шапка
    _add_paragraph(doc, f"В {d.get('court_name', 'Арбитражный суд')}", alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(doc, "", space_after=2)
    _add_paragraph(doc, "Заявитель (должник):", bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(doc, full_name, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    if d.get("inn"):
        _add_paragraph(doc, f"ИНН: {d['inn']}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)

    address_parts = [d.get("region", ""), d.get("city", ""), d.get("address", "")]
    address = ", ".join(p for p in address_parts if p)
    if address:
        _add_paragraph(doc, f"Адрес: {address}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)
    if d.get("phone"):
        _add_paragraph(doc, f"Тел.: {d['phone']}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=11)

    # Заголовок
    _add_paragraph(doc, "", space_after=12)
    _add_paragraph(doc, "ХОДАТАЙСТВО", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)
    _add_paragraph(
        doc,
        "о введении процедуры реализации имущества гражданина",
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=13,
    )
    _add_paragraph(doc, "", space_after=6)

    # Текст
    total_debt = d.get("total_debt", 0)
    monthly_income = d.get("monthly_income", 0)
    dependents = d.get("dependents_count", 0)
    subsistence = 17_733  # Прожиточный минимум 2026

    _add_paragraph(
        doc,
        f"В производстве суда находится дело о признании меня, {full_name}, "
        f"несостоятельным (банкротом). Общая сумма задолженности составляет "
        f"{_fmt_money(total_debt)} руб.",
    )

    _add_paragraph(doc, "", space_after=4)
    _add_paragraph(
        doc,
        f"Мой ежемесячный доход составляет {_fmt_money(monthly_income)} руб. "
        f"Количество лиц, находящихся на моём иждивении: {dependents}.",
    )

    income_threshold = subsistence * (1 + dependents)
    _add_paragraph(
        doc,
        f"Прожиточный минимум для трудоспособного населения составляет "
        f"{_fmt_money(subsistence)} руб. С учётом иждивенцев минимально необходимый "
        f"доход составляет {_fmt_money(income_threshold)} руб., что превышает мой "
        f"фактический доход.",
    )

    _add_paragraph(doc, "", space_after=4)
    _add_paragraph(
        doc,
        "Таким образом, у меня отсутствует источник дохода, достаточный для исполнения "
        "плана реструктуризации долгов. Введение процедуры реструктуризации долгов "
        "не приведёт к восстановлению моей платёжеспособности и лишь затянет процедуру "
        "банкротства.",
    )

    _add_paragraph(doc, "", space_after=4)
    _add_paragraph(
        doc,
        "В соответствии с п. 8 ст. 213.6 Федерального закона от 26.10.2002 N 127-ФЗ "
        "\"О несостоятельности (банкротстве)\" по результатам рассмотрения обоснованности "
        "заявления о признании гражданина банкротом арбитражный суд вправе вынести решение "
        "о признании обоснованным указанного заявления и введении реализации имущества "
        "гражданина, если гражданин не соответствует требованиям для утверждения "
        "плана реструктуризации долгов.",
        font_size=11,
    )

    # Просительная часть
    _add_paragraph(doc, "", space_after=8)
    _add_paragraph(doc, "На основании изложенного,", bold=True)
    _add_paragraph(doc, "ПРОШУ:", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(
        doc,
        "Ввести в отношении меня процедуру реализации имущества гражданина, "
        "минуя процедуру реструктуризации долгов.",
    )

    # Подпись
    _add_paragraph(doc, "", space_after=16)
    _add_paragraph(doc, 'Дата: "____" _____________ 20___ г.')
    _add_paragraph(doc, "", space_after=8)
    _add_paragraph(doc, f"Подпись: _________________ / {full_name} /")

    doc.save(str(output_path))
    return output_path


def generate_inventory(pipeline_data: dict, output_path: Path) -> Path:
    """Генерирует опись имущества в DOCX."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)

    d = pipeline_data
    full_name = " ".join(p for p in [d.get("last_name", ""), d.get("first_name", ""), d.get("middle_name", "")] if p)

    _add_paragraph(doc, "ОПИСЬ ИМУЩЕСТВА ГРАЖДАНИНА", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)
    _add_paragraph(doc, full_name, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=12)
    _add_paragraph(doc, "", space_after=8)

    properties = d.get("properties", [])
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    for i, text in enumerate(["N", "Вид имущества", "Описание", "Стоимость, руб.", "Примечание"]):
        headers[i].text = text
        for p in headers[i].paragraphs:
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(9)

    prop_type_names = {
        "apartment": "Квартира", "house": "Дом", "car": "Автомобиль",
        "land": "Земельный участок", "deposit": "Вклад/Счёт", "other": "Другое",
    }

    total_value = 0
    count = 0
    for p in properties:
        if not p.get("property_type"):
            continue
        count += 1
        row = table.add_row().cells
        row[0].text = str(count)
        row[1].text = prop_type_names.get(p["property_type"], p["property_type"])
        row[2].text = p.get("description", "")
        val = p.get("estimated_value", 0)
        row[3].text = _fmt_money(val)
        row[4].text = "Единственное жильё" if p.get("is_sole_housing") else ""
        total_value += val

    if count == 0:
        row = table.add_row().cells
        row[0].text = "—"
        row[1].text = "Имущество отсутствует"
        row[2].text = ""
        row[3].text = "0.00"
        row[4].text = ""

    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = "ИТОГО:"
    total_row[2].text = ""
    total_row[3].text = _fmt_money(total_value)
    total_row[4].text = ""

    _add_paragraph(doc, "", space_after=12)
    _add_paragraph(doc, 'Дата: "____" _____________ 20___ г.')
    _add_paragraph(doc, f"Подпись: _________________ / {full_name} /")

    doc.save(str(output_path))
    return output_path
